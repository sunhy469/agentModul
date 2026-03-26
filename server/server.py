import json
import os
import platform
import re
import shlex
import shutil
import subprocess
import time
import uuid
import xml.etree.ElementTree as ET
from pathlib import Path
from pathlib import PureWindowsPath
from typing import Any
from urllib.parse import quote_plus

import httpx
from docx import Document
from dotenv import load_dotenv
from fastmcp import FastMCP

# 初始化 MCP 服务器
mcp = FastMCP("WeatherServer")
load_dotenv()

# OpenWeather API 配置
OPENWEATHER_API_BASE = "https://api.openweathermap.org/data/2.5/weather"
API_KEY = "b9b91b1fba4fa6a2f5675288e673f9f4"
USER_AGENT = "weather-app/1.0"

# 保存目录（建议固定目录）
BASE_DIR = os.path.join(os.getcwd(), "generated_files")
os.makedirs(BASE_DIR, exist_ok=True)
SAFE_ROOT = Path(os.getenv("SAFE_WORK_ROOT", BASE_DIR)).expanduser().resolve()
PENDING_FILE_OPERATIONS: dict[str, dict[str, Any]] = {}
FEISHU_TOKEN_CACHE: dict[str, Any] = {"token": "", "expire_at": 0.0}


async def _search_arxiv(query: str, max_results: int = 5) -> list[dict[str, str]]:
    url = (
        "http://export.arxiv.org/api/query"
        f"?search_query=all:{quote_plus(query)}&start=0&max_results={max(1, min(max_results, 10))}"
    )
    async with httpx.AsyncClient(timeout=20.0) as client:
        resp = await client.get(url, headers={"User-Agent": USER_AGENT})
        resp.raise_for_status()

    root = ET.fromstring(resp.text)
    ns = {"atom": "http://www.w3.org/2005/Atom"}
    entries = root.findall("atom:entry", ns)
    results: list[dict[str, str]] = []
    for entry in entries:
        title = (entry.findtext("atom:title", default="", namespaces=ns) or "").strip().replace("\n", " ")
        summary = (entry.findtext("atom:summary", default="", namespaces=ns) or "").strip().replace("\n", " ")
        link = ""
        for ln in entry.findall("atom:link", ns):
            href = ln.attrib.get("href", "")
            rel = ln.attrib.get("rel", "")
            if href and (rel == "alternate" or "arxiv.org/abs/" in href):
                link = href
                break
        if not link:
            link = (entry.findtext("atom:id", default="", namespaces=ns) or "").strip()
        if title:
            results.append({"title": title, "summary": summary[:280], "url": link})
    return results


async def _search_duckduckgo(query: str, max_results: int = 5) -> list[dict[str, str]]:
    params = {
        "q": query,
        "format": "json",
        "no_html": "1",
        "skip_disambig": "1",
    }
    async with httpx.AsyncClient(timeout=20.0) as client:
        resp = await client.get("https://api.duckduckgo.com/", params=params, headers={"User-Agent": USER_AGENT})
        resp.raise_for_status()
        data = resp.json()

    items: list[dict[str, str]] = []
    for topic in data.get("RelatedTopics", []):
        if isinstance(topic, dict) and "Topics" in topic:
            for sub in topic.get("Topics", []):
                text = (sub.get("Text") or "").strip()
                url = (sub.get("FirstURL") or "").strip()
                if text and url:
                    items.append({"title": text, "summary": "", "url": url})
        else:
            text = (topic.get("Text") or "").strip() if isinstance(topic, dict) else ""
            url = (topic.get("FirstURL") or "").strip() if isinstance(topic, dict) else ""
            if text and url:
                items.append({"title": text, "summary": "", "url": url})
        if len(items) >= max_results:
            break

    abstract = (data.get("AbstractText") or "").strip()
    abstract_url = (data.get("AbstractURL") or "").strip()
    if abstract and abstract_url:
        items.insert(0, {"title": abstract[:120], "summary": abstract[:280], "url": abstract_url})

    return items[:max(1, min(max_results, 10))]


async def fetch_weather(city: str) -> dict[str, Any] | None:
    """
    从 OpenWeather API 获取天气信息。
    演示模式下返回模拟数据，实际使用时连接真实API
    """
    # 演示模式：返回模拟天气数据
    demo_data = {
        "name": city.title(),
        "sys": {"country": "CN"},
        "main": {
            "temp": 15,
            "humidity": 60
        },
        "wind": {"speed": 2.5},
        "weather": [{"description": "晴朗"}]
    }

    # 如果提供了真实API密钥，则尝试连接真实API
    if API_KEY != "demo-key" and API_KEY != "b9b91b1fba4fa6a2f5675288e673f9f4":
        params = {
            "q": city,
            "appid": API_KEY,
            "units": "metric",
            "lang": "zh_cn"
        }
        headers = {"User-Agent": USER_AGENT}

        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    OPENWEATHER_API_BASE,
                    params=params,
                    headers=headers,
                    timeout=30.0
                )
                response.raise_for_status()
                return response.json()

        except httpx.HTTPStatusError as e:
            return {"error": f"HTTP 错误: {e.response.status_code}"}

        except Exception as e:
            return {"error": f"请求失败: {str(e)}"}

    return demo_data


def format_weather(data: dict[str, Any] | str) -> str:
    """
    将天气数据格式化为易读文本。
    """
    # 如果传入的是字符串，则先转换为字典
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception as e:
            return f"无法解析天气数据：{e}"

    # 如果数据中包含错误信息，直接返回错误提示
    if "error" in data:
        return f"{data['error']}"

    # 提取数据时做容错处理
    city = data.get("name", "未知")
    country = data.get("sys", {}).get("country", "未知")
    temp = data.get("main", {}).get("temp", "N/A")
    humidity = data.get("main", {}).get("humidity", "N/A")
    wind_speed = data.get("wind", {}).get("speed", "N/A")

    # weather 可能为空列表，因此用 [0] 前先提供默认字典
    weather_list = data.get("weather", [{}])
    description = weather_list[0].get("description", "未知")

    return (
        f"🌍 {city}, {country}\n"
        f"🌡️ 温度: {temp}°C\n"
        f"💧 湿度: {humidity}%\n"
        f"💨 风速: {wind_speed} m/s\n"
        f"☁️ 天气: {description}\n"
    )


@mcp.tool()
def create_word_file(filename: str, content: str) -> str:
    """
    根据文件名和内容创建 Word 文件。
    :param filename : 文件名
    :param content: 内容
    """

    try:
        # 防止路径攻击
        filename = filename.replace("/", "").replace("\\", "")

        full_path = os.path.join(BASE_DIR, f"{filename}")

        doc = Document()
        doc.add_heading(filename, level=1)

        # 按行写入
        for line in content.split("\n"):
            doc.add_paragraph(line)

        doc.save(full_path)

        return f"Word 文件已创建成功：{full_path}"

    except Exception as e:
        return f"创建文件失败: {e}"


def _safe_read_file(file_path: Path, max_chars: int = 12000) -> str:
    """读取文件内容并做基础截断，避免返回过大文本。"""
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        doc = Document(str(file_path))
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        try:
            content = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return "该文件不是 UTF-8 文本，暂不支持直接解析。"
        except Exception as e:
            return f"读取文件失败: {e}"

    if not content.strip():
        return "文件内容为空。"

    if len(content) > max_chars:
        return content[:max_chars] + f"\n\n...（内容过长，已截断，原始长度 {len(content)} 字符）"

    return content


def _resolve_search_root(directory: str = "") -> Path:
    if directory.strip():
        target = Path(directory).expanduser().resolve()
    else:
        target = Path(os.getenv("LOCAL_SEARCH_DIR", BASE_DIR)).resolve()

    # 安全边界：仅允许在 SAFE_ROOT 下检索
    try:
        target.relative_to(SAFE_ROOT)
    except ValueError:
        return SAFE_ROOT

    return target


def _is_within_safe_root(target: Path) -> bool:
    try:
        target.resolve().relative_to(SAFE_ROOT)
        return True
    except Exception:
        return False


def _ensure_safe_target(file_path: str) -> tuple[Path | None, str]:
    if not file_path.strip():
        return None, "请提供文件路径。"
    target = Path(file_path).expanduser().resolve()
    if not target.exists() or not target.is_file():
        return None, f"文件不存在：{target}"
    if not _is_within_safe_root(target):
        return None, f"拒绝访问：目标超出安全目录 {SAFE_ROOT}"
    return target, ""


def _find_file_by_keyword(filename: str, directory: str = "") -> tuple[Path | None, str]:
    """按关键字查找文件并返回最佳匹配。"""
    search_root = _resolve_search_root(directory)
    if not search_root.exists() or not search_root.is_dir():
        return None, f"目录无效：{search_root}"

    keyword = filename.strip().lower()
    if not keyword:
        return None, "请提供文件名或关键字。"

    candidates = [
        p for p in search_root.rglob("*")
        if p.is_file() and keyword in p.name.lower()
    ]

    if not candidates:
        return None, f"未找到包含关键字“{filename}”的文件。搜索目录：{search_root}"

    exact = [p for p in candidates if p.name.lower() == keyword]
    target = sorted(exact or candidates, key=lambda x: (len(str(x)), str(x)))[0]
    return target, ""




@mcp.tool()
def find_and_read_local_file(filename: str, requirement: str = "", open_with_default: bool = True) -> str:
    """
    根据文件名在本地目录中查找文件，并执行“读取内容 + 默认应用打开”。
    :param filename: 待查找的文件名（支持完整文件名或关键字）
    :param requirement: 用户需求（可选），会附在返回文本中方便模型解析
    :param open_with_default: 是否按系统默认方式打开文件（默认 True）
    :return: 文件定位与内容
    """
    try:
        target, err = _find_file_by_keyword(filename)
        if err:
            return err

        if open_with_default:
            _open_with_default_app(target)

        file_content = _safe_read_file(target)
        requirement_text = requirement.strip()

        response = (
            f"已找到文件：{target}\n"
            f"文件名：{target.name}\n"
            f"文件大小：{target.stat().st_size} bytes\n"
        )

        if open_with_default:
            response += "文件已按系统默认方式打开。\n"

        if requirement_text:
            response += f"用户需求：{requirement_text}\n"

        response += f"文件内容：\n{file_content}"
        return response

    except Exception as e:
        return f"查找、读取或打开文件失败: {e}"




def _resolve_windows_app_path(app_command: str) -> str:
    """在 Windows 上尝试解析常见应用别名与安装路径。"""
    alias_map = {
        "qq": ["QQScLauncher.exe", "QQ.exe", "QQNT.exe"],
        "wechat": ["WeChat.exe"],
        "weixin": ["WeChat.exe"],
        "chrome": ["chrome.exe"],
        "edge": ["msedge.exe"],
    }

    raw = app_command.strip().strip('"')
    lower_name = raw.lower()

    # 已经是路径
    direct_path = Path(raw).expanduser()
    if direct_path.exists() and direct_path.is_file():
        return str(direct_path.resolve())

    # 先走 PATH
    which_path = shutil.which(raw)
    if which_path:
        return which_path

    candidates = []
    if lower_name in alias_map:
        candidates.extend(alias_map[lower_name])

    if lower_name.endswith('.exe'):
        candidates.append(raw)
    else:
        candidates.extend([raw, f"{raw}.exe"])

    roots = [
        Path(os.getenv("ProgramFiles", "C:/Program Files")),
        Path(os.getenv("ProgramFiles(x86)", "C:/Program Files (x86)")),
        Path(os.getenv("LOCALAPPDATA", "")),
    ]

    # 针对 QQ 的高频目录先做快速定位
    qq_fast_paths = [
        "Tencent/QQ/Bin/QQScLauncher.exe",
        "Tencent/QQ/Bin/QQ.exe",
        "Tencent/QQNT/QQ.exe",
        "Tencent/QQNT/QQNT.exe",
    ]
    if lower_name == "qq":
        for root in roots:
            for rel in qq_fast_paths:
                candidate = root / rel
                if candidate.exists():
                    return str(candidate.resolve())

    # 通用浅层搜索（限制层级，避免太慢）
    for root in roots:
        if not root.exists():
            continue
        for exe_name in candidates:
            try:
                for current_root, dirs, files in os.walk(root):
                    # 限制搜索深度
                    depth = Path(current_root).relative_to(root).parts
                    if len(depth) > 4:
                        dirs[:] = []
                        continue
                    for f in files:
                        if f.lower() == exe_name.lower():
                            return str((Path(current_root) / f).resolve())
            except Exception:
                continue

    return ""


def _build_launch_command(app_command: str, arguments: str = "") -> list[str]:
    """构造安全的启动命令，避免 shell 注入。"""
    system_name = platform.system().lower()
    raw = app_command.strip()
    if not raw:
        return []

    blocked_tokens = ["&&", "||", ";", "|", "`", "$("]
    if any(token in raw for token in blocked_tokens) or any(token in arguments for token in blocked_tokens):
        return []

    if system_name == "windows":
        executable = _resolve_windows_app_path(raw) or raw
        args = shlex.split(arguments, posix=False) if arguments.strip() else []
        return [executable, *args]

    executable = shutil.which(raw) or raw
    args = shlex.split(arguments) if arguments.strip() else []
    return [executable, *args]


@mcp.tool()
def open_local_application(app_command: str, arguments: str = "") -> str:
    """
    打开本地应用（通过命令行）。
    :param app_command: 应用启动命令，例如 qq、notepad、calc、code
    :param arguments: 可选启动参数
    """
    try:
        launch_cmd = _build_launch_command(app_command, arguments)
        if not launch_cmd:
            return "请提供应用启动命令。"

        subprocess.Popen(launch_cmd, shell=False)
        return f"已尝试启动应用：{' '.join(launch_cmd)}"
    except FileNotFoundError:
        return (
            "未找到可执行程序。可尝试：\n"
            "1) 直接传完整 exe 路径；\n"
            "2) 使用常见别名（如 qq/wechat/chrome）；\n"
            "3) 将应用目录加入系统 PATH。"
        )
    except Exception as e:
        return f"启动应用失败: {e}"


def _open_with_default_app(target: Path) -> None:
    """按系统默认方式打开文件。"""
    system_name = platform.system().lower()
    if system_name == "windows":
        os.startfile(str(target))
    elif system_name == "darwin":
        subprocess.Popen(["open", str(target)])
    else:
        subprocess.Popen(["xdg-open", str(target)])


def _is_windows_process_running(process_names: list[str]) -> bool:
    if platform.system().lower() != "windows":
        return False
    try:
        result = subprocess.run(
            ["tasklist"],
            capture_output=True,
            text=True,
            timeout=3,
            check=False,
        )
        output = (result.stdout or "").lower()
        return any(name.lower() in output for name in process_names)
    except Exception:
        return False


def _wait_for_qq_window(timeout_seconds: float = 15.0) -> tuple[tuple[int, int, int, int] | None, str]:
    """等待 QQ 窗口可操作后返回窗口位置。"""
    deadline = time.time() + max(1.0, timeout_seconds)
    while time.time() < deadline:
        try:
            import pygetwindow as gw
            windows = gw.getWindowsWithTitle("QQ")
            for w in windows:
                if w.width > 200 and w.height > 200:
                    if w.isMinimized:
                        w.restore()
                    try:
                        w.activate()
                    except Exception:
                        pass
                    return (w.left, w.top, w.width, w.height), ""
        except Exception:
            return None, "缺少 pygetwindow 或窗口探测失败，无法确认 QQ 是否完全就绪。"
        time.sleep(0.25)
    return None, f"等待 QQ 窗口超时（{timeout_seconds:.1f}s），任务未完成。"


@mcp.tool()
def read_file(file_path: str, open_with_default: bool = True) -> str:
    """
    读取并（可选）按系统默认方式打开文件。
    :param file_path: 文件完整路径
    :param open_with_default: 是否按默认应用打开（默认 True）
    """
    try:
        if not file_path.strip():
            return "请提供文件路径。"

        target = Path(file_path).expanduser().resolve()
        if not target.exists() or not target.is_file():
            return f"文件不存在：{target}"
        if not _is_within_safe_root(target):
            return f"拒绝访问：目标超出安全目录 {SAFE_ROOT}"

        if open_with_default:
            _open_with_default_app(target)
            return f"已按系统默认方式打开文件：{target}"

        # 兼容部分请求：如果不打开则返回内容摘要
        content = _safe_read_file(target)
        return f"文件路径：{target}\n文件内容：\n{content}"
    except Exception as e:
        return f"读取或打开文件失败: {e}"


@mcp.tool()
def open_path_in_file_manager(target_path: str = "") -> str:
    """
    在系统文件管理器中打开目录或文件。
    :param target_path: 目标路径，留空则打开 BASE_DIR
    """
    try:
        target = Path(target_path).expanduser().resolve() if target_path.strip() else Path(BASE_DIR).resolve()

        if not target.exists():
            return f"目标路径不存在：{target}"

        system_name = platform.system().lower()
        if system_name == "windows":
            os.startfile(str(target))
        elif system_name == "darwin":
            subprocess.Popen(["open", str(target)])
        else:
            subprocess.Popen(["xdg-open", str(target)])

        return f"已在文件管理器打开：{target}"
    except Exception as e:
        return f"打开路径失败: {e}"


@mcp.tool()
def list_local_files(directory: str = "", keyword: str = "", limit: int = 20) -> str:
    """
    列出本地文件，便于后续让模型选择要读取的文件。
    :param directory: 搜索目录，默认 LOCAL_SEARCH_DIR 或 BASE_DIR
    :param keyword: 文件名过滤关键字
    :param limit: 最多返回条数
    """
    try:
        search_root = _resolve_search_root(directory)
        if not search_root.exists() or not search_root.is_dir():
            return f"目录无效：{search_root}"

        kw = keyword.strip().lower()
        matches = []
        for p in search_root.rglob("*"):
            if not p.is_file():
                continue
            if kw and kw not in p.name.lower():
                continue
            matches.append(p)

        if not matches:
            return f"目录中未找到匹配文件。目录：{search_root}"

        limit = max(1, min(limit, 100))
        selected = sorted(matches, key=lambda x: str(x))[:limit]

        lines = [f"搜索目录：{search_root}", f"匹配数量：{len(matches)}（展示前 {len(selected)} 条）"]
        for idx, item in enumerate(selected, 1):
            lines.append(f"{idx}. {item.name} | {item}")

        return "\n".join(lines)
    except Exception as e:
        return f"列出文件失败: {e}"


@mcp.tool()
def create_pending_file_operation(
        operation: str,
        file_path: str,
        new_content: str = "",
        reason: str = "",
) -> str:
    """
    创建待确认文件操作（delete/modify），用于双重安全确认。
    """
    op = operation.strip().lower()
    if op not in {"delete", "modify"}:
        return "operation 仅支持 delete 或 modify。"

    target, err = _ensure_safe_target(file_path)
    if err:
        return err

    if op == "modify" and not new_content.strip():
        return "modify 操作必须提供 new_content。"

    operation_id = uuid.uuid4().hex[:12]
    PENDING_FILE_OPERATIONS[operation_id] = {
        "operation": op,
        "target": str(target),
        "new_content": new_content,
        "reason": reason.strip(),
        "created_at": time.time(),
    }
    return (
        f"已创建待确认操作，operation_id={operation_id}\n"
        f"类型: {op}\n目标: {target}\n"
        "请调用 confirm_file_operation(operation_id, confirm=True) 执行；"
        "若取消请传 confirm=False。"
    )


@mcp.tool()
def confirm_file_operation(operation_id: str, confirm: bool) -> str:
    """
    执行或取消待确认文件操作，防止误删误改。
    """
    payload = PENDING_FILE_OPERATIONS.pop(operation_id, None)
    if not payload:
        return f"未找到待确认操作：{operation_id}"

    if not confirm:
        return f"已取消操作：{operation_id}"

    target = Path(payload["target"]).expanduser().resolve()
    if not _is_within_safe_root(target):
        return f"拒绝执行：目标超出安全目录 {SAFE_ROOT}"

    if payload["operation"] == "delete":
        target.unlink(missing_ok=False)
        return f"已删除文件：{target}"

    target.write_text(payload["new_content"], encoding="utf-8")
    return f"已更新文件：{target}（UTF-8 覆盖写入）"


@mcp.tool()
def get_pending_operations() -> str:
    """查看尚未确认的文件操作。"""
    if not PENDING_FILE_OPERATIONS:
        return "当前没有待确认文件操作。"

    lines = [f"待确认操作数量: {len(PENDING_FILE_OPERATIONS)}"]
    for op_id, payload in PENDING_FILE_OPERATIONS.items():
        lines.append(f"- {op_id} | {payload['operation']} | {payload['target']}")
    return "\n".join(lines)


@mcp.tool()
def analyze_time_series(values: list[float], horizon: int = 3) -> str:
    """
    时间序列快速分析：给出均值、趋势与线性外推预测。
    """
    if not values or len(values) < 3:
        return "请至少提供 3 个数值。"
    horizon = max(1, min(30, int(horizon)))
    n = len(values)
    avg = sum(values) / n
    trend = (values[-1] - values[0]) / (n - 1)
    forecast = [round(values[-1] + trend * (i + 1), 4) for i in range(horizon)]
    return (
        f"样本数: {n}\n均值: {avg:.4f}\n"
        f"线性趋势(每步): {trend:.4f}\n"
        f"未来 {horizon} 步预测: {forecast}"
    )


def _resolve_feishu_robot_url(robot_name: str = "default") -> str:
    normalized = (robot_name or "default").strip().upper().replace("-", "_")
    if normalized == "DEFAULT":
        return os.getenv("FEISHU_ROBOT_URL", "").strip() or os.getenv("FEISHU_WEBHOOK_URL", "").strip()
    return (
            os.getenv(f"FEISHU_ROBOT_{normalized}_URL", "").strip()
            or os.getenv("FEISHU_ROBOT_URL", "").strip()
            or os.getenv("FEISHU_WEBHOOK_URL", "").strip()
    )


async def _get_feishu_tenant_access_token() -> tuple[str, str]:
    now = time.time()
    cached_token = str(FEISHU_TOKEN_CACHE.get("token", "")).strip()
    cached_expire_at = float(FEISHU_TOKEN_CACHE.get("expire_at", 0.0))
    if cached_token and now < (cached_expire_at - 60):
        return cached_token, ""

    app_id = os.getenv("FEISHU_APP_ID", "").strip()
    app_secret = os.getenv("FEISHU_APP_SECRET", "").strip()
    if not app_id or not app_secret:
        return "", "缺少 FEISHU_APP_ID 或 FEISHU_APP_SECRET 环境变量。"

    token_url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
    payload = {
        "app_id": app_id,
        "app_secret": app_secret,
    }

    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            resp = await client.post(token_url, json=payload)
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return "", f"获取 tenant_access_token 失败: {exc}"

    if data.get("code") != 0:
        return "", f"获取 tenant_access_token 失败: code={data.get('code')}, msg={data.get('msg')}"

    token = data.get("tenant_access_token", "")
    if not token:
        return "", "获取 tenant_access_token 成功但响应中无 tenant_access_token。"

    expires_in = int(data.get("expire", 7200) or 7200)
    FEISHU_TOKEN_CACHE["token"] = token
    FEISHU_TOKEN_CACHE["expire_at"] = now + max(60, expires_in)
    return token, ""


def _resolve_image_file(image_path: str) -> Path | None:
    """尽量把“图片名”解析为可读取的本地文件路径。"""
    raw = (image_path or "").strip().strip('"').strip("'")
    if not raw:
        return None

    # 先尝试输入本身（支持绝对路径）
    direct_candidates = [raw]
    if "\\" in raw:
        direct_candidates.append(raw.replace("\\", "/"))
    for item in direct_candidates:
        p = Path(item).expanduser()
        if p.exists() and p.is_file():
            return p.resolve()

    # 若只传了文件名或 Windows 风格路径，统一提取文件名
    base_name = PureWindowsPath(raw).name if ("\\" in raw or ":" in raw) else Path(raw).name
    candidates = [
        Path.cwd() / base_name,
        Path(BASE_DIR) / base_name,
        Path(os.getenv("LOCAL_SEARCH_DIR", BASE_DIR)).expanduser().resolve() / base_name,
        Path(os.getenv("SAFE_WORK_ROOT", BASE_DIR)).expanduser().resolve() / base_name,
        ]
    for c in candidates:
        if c.exists() and c.is_file():
            return c.resolve()

    # 深度检索在大目录（如 D:\）会很慢，默认关闭；需要时可显式开启。
    deep_search = os.getenv("ENABLE_DEEP_IMAGE_SEARCH", "0").strip().lower() in {"1", "true", "yes"}
    if deep_search:
        try:
            for fp in SAFE_ROOT.rglob(base_name):
                if fp.is_file():
                    return fp.resolve()
        except Exception:
            return None
    return None


async def _upload_feishu_image(
        image_path: str,
        tenant_access_token: str,
) -> tuple[str, str]:
    if not image_path.strip():
        return "", "image_path 不能为空。"
    if not tenant_access_token.strip():
        return "", "发送图片消息需要提供 tenant_access_token。"

    upload_url = "https://open.feishu.cn/open-apis/im/v1/images"
    headers = {"Authorization": f"Bearer {tenant_access_token.strip()}"}
    data = {"image_type": "message"}

    try:
        file_path = _resolve_image_file(image_path)
        if not file_path:
            return "", f"图片文件不存在或无法定位：{image_path}"
        with file_path.open("rb") as fp:
            files = {"image": (file_path.name, fp, "application/octet-stream")}
            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.post(upload_url, headers=headers, data=data, files=files)
        resp.raise_for_status()
        payload = resp.json()
    except Exception as exc:
        return "", f"上传图片到飞书失败: {exc}"

    if payload.get("code") != 0:
        return "", f"上传图片到飞书失败: code={payload.get('code')}, msg={payload.get('msg')}"

    image_key = payload.get("data", {}).get("image_key", "")
    if not image_key:
        return "", "上传图片成功但未返回 image_key。"
    return image_key, ""


async def _send_feishu_message(
        message: str,
        link_text: str = "",
        link_href: str = "",
        image_path: str = "",
) -> str:
    if not message.strip():
        return "message 不能为空。"
    robot_url = os.getenv("FEISHU_WEBHOOK_URL");

    if not robot_url:
        return "未配置 FEISHU_WEBHOOK_URL。"

    if image_path.strip():
        token, token_err = await _get_feishu_tenant_access_token()
        if token_err:
            return token_err

        image_key, err = await _upload_feishu_image(
            image_path=image_path,
            tenant_access_token=token,
        )
        if err:
            return err
        payload: dict[str, Any] = {
            "msg_type": "image",
            "content": {"image_key": image_key}
        }
    elif link_text.strip() and link_href.strip():
        nodes: list[dict[str, Any]] = [{"tag": "text", "text": message}]
        nodes.append({"tag": "a", "text": link_text.strip(), "href": link_href.strip()})
        payload = {
            "msg_type": "post",
            "content": {
                "post": {
                    "zh_cn": {
                        "title": "",
                        "content": [nodes]
                    }
                }
            },
        }
    else:
        payload = {"msg_type": "text", "content": {"text": message}}

    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.post(robot_url, json=payload)
            if resp.is_success:
                return f"飞书机器人发送完成，HTTP {resp.status_code}。"
            return f"飞书机器人发送失败，HTTP {resp.status_code}，响应：{resp.text}"
    except Exception as exc:
        return f"飞书机器人发送失败: {exc}"


@mcp.tool()
async def send_feishu_robot_message(
        message: str,
        link_text: str = "",
        link_href: str = "",
        image_path: str = "",
) -> str:
    """
    使用飞书机器人发送消息（支持 text、带超链接的 post、image）。
    :param message: 文本内容（必填）
    :param link_text: 超链接显示文字（可选）
    :param link_href: 超链接地址（可选）
    :param image_path: 本地图片路径或图片文件名（可选，存在时发送 image 消息）
    """
    return await _send_feishu_message(
        message=message,
        link_text=link_text,
        link_href=link_href,
        image_path=image_path,
    )


@mcp.tool()
def summarize_local_capabilities() -> str:
    """
    返回当前可用于本地自动化的能力说明，便于多工具编排。
    """
    return (
        "当前本地自动化能力：\n"
        "1) 文件系统：查找/读取/列举/待确认修改删除。\n"
        "2) 本地应用：按命令启动应用（可用于拉起 QQ、飞书、浏览器）。\n"
        "3) 通知消息：支持飞书机器人发送。\n"
        "4) 在线检索：支持通用网页检索与 arXiv 论文检索。\n"
        "5) 桌面 IM：支持自动粘贴并发送消息（依赖 pyautogui + pyperclip）。\n"
        "6) 文档处理：可创建 Word 文件并导出。\n"
        "7) 时间序列：基础统计与趋势预测。"
    )


@mcp.tool()
async def search_web(query: str, max_results: int = 5, source: str = "auto") -> str:
    """
    在线检索工具。
    - source=auto: 自动判断，学术相关优先 arXiv
    - source=arxiv: 仅检索 arXiv
    - source=web: 通用网页检索（DuckDuckGo API）
    """
    q = (query or "").strip()
    if not q:
        return "请提供 query。"

    source_norm = (source or "auto").strip().lower()
    if source_norm not in {"auto", "arxiv", "web"}:
        source_norm = "auto"

    is_academic_query = any(k in q.lower() for k in ["论文", "文献", "arxiv", "paper", "scholar", "attention", "k-means"])
    if source_norm == "auto":
        source_norm = "arxiv" if is_academic_query else "web"

    try:
        results = await (_search_arxiv(q, max_results) if source_norm == "arxiv" else _search_duckduckgo(q, max_results))
    except Exception as exc:
        return f"在线检索失败: {exc}"

    if not results:
        return f"未检索到结果。query={q}，source={source_norm}"

    lines = [f"检索来源: {source_norm}", f"query: {q}", f"结果数: {len(results)}"]
    for i, item in enumerate(results, start=1):
        title = item.get("title", "")
        summary = item.get("summary", "")
        url = item.get("url", "")
        lines.append(f"{i}. {title}")
        if summary:
            lines.append(f"   摘要: {summary}")
        if url:
            lines.append(f"   链接: {url}")
    return "\n".join(lines)


@mcp.tool()
def send_desktop_message(
        app_command: str,
        message: str,
        press_enter: bool = True,
        warmup_seconds: float = 1.5,
        focus_input_click: bool = True,
        conversation_name: str = "",
        relaunch_if_running: bool = False,
        wait_ready_seconds: float = 15.0,
        post_action_wait_seconds: float = 0.8,
) -> str:
    """
    桌面 IM 自动发送消息（QQ/飞书/企业微信等）。
    说明：依赖 pyautogui + pyperclip，且需要桌面会话可交互。
    """
    if not app_command.strip():
        return "请提供 app_command（如 qq / feishu / wechat）。"
    if not message.strip():
        return "请提供 message 文本。"

    try:
        import pyautogui
        import pyperclip
    except Exception:
        return "缺少依赖：请安装 pyautogui 与 pyperclip 后重试。"

    running = _is_windows_process_running(["QQ.exe", "QQNT.exe", "QQScLauncher.exe"]) if app_command.lower() == "qq" else False
    if running and not relaunch_if_running:
        launch = "检测到 QQ 已在运行，跳过重复启动。"
    else:
        launch = open_local_application(app_command, "")
    time.sleep(max(0.5, min(warmup_seconds, 10.0)))

    qq_window_box: tuple[int, int, int, int] | None = None
    steps: list[str] = [launch]
    if app_command.lower() == "qq":
        qq_window_box, qq_err = _wait_for_qq_window(timeout_seconds=wait_ready_seconds)
        if qq_err:
            return f"{launch}\n{qq_err}"
        steps.append("QQ 窗口已就绪。")

    # 若提供会话名，先在 IM 中搜索并进入对应会话
    if conversation_name.strip():
        if app_command.lower() == "qq" and qq_window_box:
            left, top, width, _ = qq_window_box
            # 点击 QQ 搜索框
            pyautogui.click(left + int(width * 0.24), top + 36)
            time.sleep(0.1)
            pyautogui.hotkey("ctrl", "a")
            pyautogui.press("backspace")
            pyautogui.write(conversation_name.strip(), interval=0.03)
            time.sleep(0.2)
            pyautogui.press("enter")
            time.sleep(0.3)
            # 点击第一条会话，确保真正进入聊天
            pyautogui.click(left + int(width * 0.24), top + 126)
            time.sleep(0.2)
            steps.append(f"已执行 QQ 搜索并尝试进入会话：{conversation_name.strip()}")
        elif platform.system().lower() == "darwin":
            pyperclip.copy(conversation_name.strip())
            pyautogui.hotkey("command", "f")
            time.sleep(0.1)
            pyautogui.hotkey("command", "v")
        else:
            pyperclip.copy(conversation_name.strip())
            pyautogui.hotkey("ctrl", "f")
            time.sleep(0.1)
            pyautogui.hotkey("ctrl", "v")
        time.sleep(0.2)
        pyautogui.press("enter")
        time.sleep(0.4)

    # 尝试将焦点切到聊天输入框（对 QQ/飞书等桌面 IM 更稳定）
    if focus_input_click:
        try:
            if app_command.lower() == "qq" and qq_window_box:
                left, top, width, height = qq_window_box
                pyautogui.click(left + int(width * 0.70), top + int(height * 0.86))
            else:
                screen_w, screen_h = pyautogui.size()
                pyautogui.click(int(screen_w * 0.72), int(screen_h * 0.92))
            time.sleep(0.2)
        except Exception:
            pass

    # 优先采用“键盘输入”，仅在非 ASCII 文本时回退剪贴板粘贴
    if message.isascii():
        pyautogui.write(message, interval=0.01)
    elif platform.system().lower() == "darwin":
        pyperclip.copy(message)
        pyautogui.hotkey("command", "v")
    else:
        pyperclip.copy(message)
        pyautogui.hotkey("ctrl", "v")
        # Windows 下部分场景 Ctrl+V 不生效，补一个 Shift+Insert 兜底
        time.sleep(0.1)
        pyautogui.hotkey("shift", "insert")
    if press_enter:
        pyautogui.press("enter")
        time.sleep(0.05)
        pyautogui.press("enter")
        steps.append("已执行回车发送。")

    time.sleep(max(0.0, min(post_action_wait_seconds, 5.0)))
    steps.append("桌面自动化步骤执行完毕。")

    return "\n".join(steps)


def _detect_channel_from_request(request: str) -> str:
    text = (request or "").lower()
    if any(k in text for k in ["飞书", "feishu", "lark"]):
        return "feishu"
    if any(k in text for k in ["qq", "企鹅"]):
        return "qq"
    if any(k in text for k in ["facebook", "messenger", "meta"]):
        return "unsupported"
    return "auto"


def _extract_qq_target_and_message(request: str, fallback_message: str = "") -> tuple[str, str]:
    text = (request or "").strip()
    if not text:
        return "", fallback_message.strip()

    # 例：在qq找一下bamboo这个人，和他说，让他明天下午来找我
    m = re.search(r"(?:找(?:一下)?)([^，。,.\s]+)(?:这个人)?(?:，|,)?(?:和他说|并?告诉他|对他说|说)[:：，,]?(.*)", text)
    if m:
        name = (m.group(1) or "").strip()
        content = (m.group(2) or "").strip() or fallback_message.strip()
        return name, content
    return "", fallback_message.strip() or text


@mcp.tool()
async def send_message_by_request(
        request: str,
        message: str = "",
        preferred_channel: str = "auto",
        qq_app_command: str = "qq",
        auto_send: bool = True,
) -> str:
    """
    根据请求自动选择飞书或 QQ 发送消息。
    - 飞书：调用飞书机器人发送
    - QQ：调用桌面自动发送
    """
    channel = preferred_channel.strip().lower()
    if channel == "auto":
        channel = _detect_channel_from_request(request)
    final_message = message.strip() or request.strip()
    if not final_message:
        return "消息内容为空，请提供 request 或 message。"

    if channel == "unsupported":
        return "检测到目标为 Facebook/Messenger，当前仅支持飞书机器人与 QQ 桌面发送。"

    if channel in {"feishu", "lark"}:
        return await send_feishu_robot_message(message=final_message)
    if channel == "qq":
        target_name, qq_message = _extract_qq_target_and_message(request, fallback_message=message.strip() or final_message)
        return send_desktop_message(
            app_command=qq_app_command,
            message=qq_message or final_message,
            press_enter=auto_send,
            conversation_name=target_name,
        )

    return (
        "未识别发送渠道。请在请求中明确写“飞书”或“QQ”，"
        "或传 preferred_channel='feishu' / 'qq'。"
    )




@mcp.tool()
def execute_complex_instruction(
        instruction: str = "",
        file_keyword: str = "",
        directory: str = "",
        open_file: bool = True,
        read_content: bool = True,
        requirement: str = "",
        app_command: str = "",
        app_arguments: str = "",
        export_word_name: str = "",
        export_word_content: str = "",
) -> str:
    """
    执行复杂指令（结构化编排版）。

    设计目标：由大模型先完成语义理解，将意图映射为参数；本工具只做稳定、可控的步骤执行，
    避免在服务端做关键词匹配导致“看起来不智能”。

    参数说明：
    - instruction: 原始用户指令（用于记录）
    - file_keyword: 需要查找的文件关键字（为空则跳过文件步骤）
    - directory: 可选搜索目录
    - open_file/read_content/requirement: 文件处理行为控制
    - app_command/app_arguments: 应用启动步骤
    - export_word_name/export_word_content: 文档导出步骤
    """
    try:
        begin = time.perf_counter()
        steps: list[str] = []
        timings: list[str] = []
        if instruction.strip():
            steps.append(f"复杂指令：{instruction.strip()}")

        content_cache = ""

        # 1) 文件查找/读取/打开
        if file_keyword.strip():
            t0 = time.perf_counter()
            target, err = _find_file_by_keyword(file_keyword, directory)
            if err:
                steps.append(f"[文件步骤失败] {err}")
            else:
                steps.append(f"[文件步骤] 已定位文件：{target}")

                if open_file:
                    _open_with_default_app(target)
                    steps.append(f"[文件步骤] 已按默认应用打开：{target}")

                if read_content:
                    content_cache = _safe_read_file(target)
                    steps.append(f"[文件步骤] 已读取内容（长度 {len(content_cache)}）")

                if requirement.strip():
                    steps.append(f"[文件步骤] 解析需求：{requirement.strip()}")
            timings.append(f"[耗时] 文件阶段: {(time.perf_counter() - t0) * 1000:.2f} ms")

        # 2) 打开应用
        if app_command.strip():
            t1 = time.perf_counter()
            app_result = open_local_application(app_command, app_arguments)
            steps.append(f"[应用步骤] {app_result}")
            timings.append(f"[耗时] 应用阶段: {(time.perf_counter() - t1) * 1000:.2f} ms")

        # 3) 导出 Word
        final_export_name = export_word_name.strip()
        final_export_content = export_word_content.strip()
        if final_export_name:
            t3 = time.perf_counter()
            if not final_export_content:
                final_export_content = content_cache[:3000] if content_cache else "（未提供导出内容）"
            export_result = create_word_file(final_export_name, final_export_content)
            steps.append(f"[导出步骤] {export_result}")
            timings.append(f"[耗时] 导出阶段: {(time.perf_counter() - t3) * 1000:.2f} ms")

        if len(steps) == 0:
            return "未接收到可执行的复杂步骤参数，请由模型先解析意图后传入结构化参数。"

        steps.extend(timings)
        steps.append(f"[性能汇总] 总耗时: {(time.perf_counter() - begin) * 1000:.2f} ms")
        return "\n".join(steps)
    except Exception as e:
        return f"复杂指令执行失败: {e}"

@mcp.tool()
async def query_weather(city: str) -> str:
    """
    输入指定城市的英文名称，返回今日天气查询结果。
    :param city: 城市名称
    :return: 格式化后的天气信息
    """
    data = await fetch_weather(city)
    return format_weather(data)


@mcp.tool()
async def get_weather_forecast(city: str, days: int = 3) -> str:
    """
    获取指定城市未来几天的天气预报（演示功能）
    :param city: 城市名称
    :param days: 预报天数（默认3天）
    :return: 天气预报信息
    """
    # 演示模式的天气预报
    forecast_data = {
        "city": city.title(),
        "forecast": [
            {"day": f"{i + 1}天", "temp": f"{20 + i}℃", "weather": "晴朗"}
            for i in range(days)
        ]
    }
    result = f"{forecast_data['city']} 未来{days}天天气预报：\n"
    for day_info in forecast_data['forecast']:
        result += f"{day_info['day']}: {day_info['temp']} - {day_info['weather']}\n"

    return result


if __name__ == "__main__":
    # 以标准 I/O 方式运行 MCP 服务器
    mcp.run(transport='stdio')
